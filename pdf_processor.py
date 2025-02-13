import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import fitz  # PyMuPDF
import docx
import re
import os
import sys

# --------------------------------------------------------------------------
# Helper: Get resource path (works for PyInstaller bundle)
# --------------------------------------------------------------------------
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


# --------------------------------------------------------------------------
# 1) Helper: place_page_full => top-aligned, minimal whitespace
# --------------------------------------------------------------------------
def place_page_full(doc_out_page, doc_in, in_page_index, target_rect):
    """
    Place page `in_page_index` from `doc_in` into `doc_out_page`
    so it fills `target_rect` from the TOP, preserving aspect ratio.
    """
    in_page = doc_in[in_page_index]
    in_rect = in_page.rect
    scale = min(
        target_rect.width / in_rect.width,
        target_rect.height / in_rect.height
    )
    new_w = in_rect.width * scale
    new_h = in_rect.height * scale

    left = target_rect.x0 + (target_rect.width - new_w) / 2
    top  = target_rect.y0
    show_rect = fitz.Rect(left, top, left + new_w, top + new_h)

    # This call can raise ValueError if page is considered "empty" by PyMuPDF
    doc_out_page.show_pdf_page(
        show_rect, doc_in, in_page_index, keep_proportion=False
    )


def create_two_up_pdf(input_pdf, output_pdf):
    """
    Produce a 2-up PDF in A4 LANDSCAPE (842 wide x 595 high).
    Left half => page i, right => page i+1.
    """
    doc_in = fitz.open(input_pdf)
    doc_out = fitz.open()
    num_pages = doc_in.page_count

    if num_pages == 0:
        # No pages to process
        doc_in.close()
        return

    # A4 Landscape
    page_width, page_height = 842, 595

    for i in range(0, num_pages, 2):
        new_page = doc_out.new_page(width=page_width, height=page_height)
        left_rect = fitz.Rect(0, 0, page_width / 2, page_height)
        right_rect = fitz.Rect(page_width / 2, 0, page_width, page_height)

        # Left page
        try:
            place_page_full(new_page, doc_in, i, left_rect)
        except ValueError:
            # Skip if "nothing to show"
            pass

        # Right page
        if i + 1 < num_pages:
            try:
                place_page_full(new_page, doc_in, i + 1, right_rect)
            except ValueError:
                pass

    doc_out.save(output_pdf)
    doc_in.close()
    doc_out.close()


def create_six_page_thumbnail_pdf(input_pdf, output_pdf):
    """
    Multi-page “thumbnail” PDF in a 3x2 grid, skipping the first (blank) page.
    Now using A4 PORTRAIT (595 wide x 842 high).
    """
    doc_in = fitz.open(input_pdf)
    doc_out = fitz.open()
    num_pages = doc_in.page_count

    if num_pages < 2:
        # If there's only the inserted blank page or zero pages, nothing to generate
        doc_in.close()
        return

    # We'll skip the first inserted page (index=0) by default:
    page_indexes = list(range(1, num_pages))

    columns, rows = 3, 2
    thumbs_per_page = columns * rows

    # A4 Portrait
    page_width, page_height = 595, 842
    slot_w = page_width / columns
    slot_h = page_height / rows

    for start_idx in range(0, len(page_indexes), thumbs_per_page):
        chunk = page_indexes[start_idx : start_idx + thumbs_per_page]
        new_page = doc_out.new_page(width=page_width, height=page_height)

        for i, src_page_index in enumerate(chunk):
            row = i // columns
            col = i % columns
            x0 = col * slot_w
            y0 = row * slot_h
            x1 = x0 + slot_w
            y1 = y0 + slot_h
            target_rect = fitz.Rect(x0, y0, x1, y1)

            try:
                place_page_full(new_page, doc_in, src_page_index, target_rect)
            except ValueError:
                # Skip if page is empty
                pass

    doc_out.save(output_pdf)
    doc_in.close()
    doc_out.close()


# --------------------------------------------------------------------------
# 2) Main PDF Processing
# --------------------------------------------------------------------------
def process_pdf_file(pdf_path,
                     output_pdf_path,
                     word_output_path,
                     personalization_word_path,
                     stamp_images,
                     watermark_text,
                     options):
    """
    Steps (only performed if corresponding checkboxes are True in 'options'):

    - Insert blank page (always, to keep consistent indexing)
    - Crop pages (if options['crop_pages'])
    - Highlight EXACT "gift message included" (if options['highlight_gmi'])
    - Highlight entire gift message snippet (if options['highlight_gift_snippet'])
    - Extract gift message & personalization text to Word docs (if options['extract_text'])
    - Highlight personalization lines (if options['highlight_personalization'])
    - Highlight quantity >=2 (if options['highlight_quantity'])
    - Add stamps (if options['add_stamps'])
    - Insert watermark text (if options['apply_watermark'])
    - Save final PDF
    """
    try:
        # Open input PDF and insert blank page at the beginning
        doc = fitz.open(pdf_path)
        doc.insert_page(0)
        total_pages = doc.page_count

        # Colors
        purple_color = (0.85, 0.6, 0.85)
        yellow_color = (1.0, 1.0, 0.0)

        # EXACT phrase => highlight purple
        highlight_only_phrases = ["gift message included"]

        # Trigger phrases to find Gift Messages
        gift_trigger_phrases = ["gift message"]

        personalization_marker = "personalization:"
        item_descriptors = ["word or message:", "word or mssg:", "morse code:"]

        # Prepare docx only if extracting is enabled
        if options['extract_text']:
            gift_doc = docx.Document()
            gift_doc.add_heading("Extracted Gift Messages", level=1)

            personalization_doc = docx.Document()
            personalization_doc.add_heading("Extracted Personalizations", level=1)
        else:
            gift_doc = None
            personalization_doc = None

        # We'll track gift messages so we don't duplicate in the doc
        gift_messages_per_page = {}

        # If you want to exclude certain areas from search, define rectangles in inches
        excluded_zones_inch = [
            (0.4313, 2.3914, 10.544, 1.7479),
            (2.3784, 8.3065, 9.743, 7.1824),
        ]

        def inches_to_points(inches):
            return inches * 72.0

        def inches_to_rect(page, left_in, right_in, top_in, bottom_in):
            page_h = page.rect.height
            x0 = inches_to_points(left_in)
            x1 = inches_to_points(right_in)
            y0 = page_h - inches_to_points(top_in)
            y1 = page_h - inches_to_points(bottom_in)
            return fitz.Rect(min(x0,x1), min(y0,y1), max(x0,x1), max(y0,y1))

        def get_excluded_rects(page):
            rects = []
            for (lft, rgt, top, bot) in excluded_zones_inch:
                rects.append(inches_to_rect(page, lft, rgt, top, bot))
            return rects

        def get_lines(page):
            """
            Returns a list of (line_text, line_rect, word_list),
            where word_list = [(x0,y0,x1,y1, text_of_word), ...]
            """
            result = []
            p_dict = page.get_text("dict")
            if not p_dict:
                return result

            all_words = page.get_text("words")  # for building line word lists

            for block in p_dict.get("blocks", []):
                for line in block.get("lines", []):
                    line_text = ""
                    line_rect = None
                    for span in line.get("spans", []):
                        span_rect = fitz.Rect(span["bbox"])
                        if line_rect is None:
                            line_rect = span_rect
                        else:
                            line_rect |= span_rect
                        line_text += span.get("text", "")
                    line_text = line_text.strip()
                    if not line_text or not line_rect:
                        continue

                    # Build the word_list for that line:
                    word_list = []
                    for w in all_words:
                        wx0, wy0, wx1, wy1, wtext = w[:5]
                        w_rect = fitz.Rect(wx0, wy0, wx1, wy1)
                        # If the word rect intersects the line rect, consider it part of the line
                        if w_rect.intersects(line_rect):
                            word_list.append((wx0, wy0, wx1, wy1, wtext))
                    # Sort them in reading order
                    word_list.sort(key=lambda x: (round(x[1], 1), x[0]))

                    result.append((line_text, line_rect, word_list))
            return result

        # Common crop rectangle
        crop_rect = fitz.Rect(0, 105, 612, 792)

        # We'll also store "ship to" info for watermark text on page0
        ship_to_info = []

        # ----------------------------------------------
        # Process each page (skip the new blank at index=0)
        # ----------------------------------------------
        for page_index in range(1, total_pages):
            page = doc[page_index]

            # (A) Crop pages if checkbox is selected
            if options['crop_pages']:
                page.set_mediabox(crop_rect)

            # Prepare to store gift messages found on this page
            gift_messages_per_page[page_index] = set()

            # (B) Highlight EXACT "gift message included"
            if options['highlight_gmi']:
                for phrase in highlight_only_phrases:
                    hits = page.search_for(phrase, flags=1)  # case-insensitive
                    for inst in hits:
                        hl = page.add_highlight_annot(inst)
                        hl.set_colors(stroke=purple_color)
                        hl.update()

            # (C) Highlight entire "gift message" snippet & extract text
            do_snippet_highlight = options['highlight_gift_snippet']
            do_text_extraction = options['extract_text']

            if do_snippet_highlight or do_text_extraction:
                for phrase in gift_trigger_phrases:
                    hits = page.search_for(phrase, flags=1)
                    if not hits:
                        continue

                    for inst in hits:
                        # Expand rectangle around "gift message"
                        expand_left  = 36
                        expand_right = 36
                        expand_up    = 72
                        expand_down  = 216
                        expanded_rect = fitz.Rect(
                            inst.x0 - expand_left,
                            inst.y0 - expand_down,
                            inst.x1 + expand_right,
                            inst.y1 + expand_up
                        )

                        # gather text blocks that intersect expanded_rect
                        blocks = page.get_text("blocks")
                        excluded_rects = get_excluded_rects(page)
                        snippet_blocks = []
                        for b in blocks:
                            bx0, by0, bx1, by1, btext, *rest = b
                            block_rect = fitz.Rect(bx0, by0, bx1, by1)
                            if not block_rect.intersects(expanded_rect):
                                continue
                            # skip blocks that intersect excluded rect
                            if any(block_rect.intersects(exr) for exr in excluded_rects):
                                continue
                            snippet_blocks.append((btext.strip(), block_rect))

                        # Build big snippet text
                        snippet_text = "\n".join([nb[0] for nb in snippet_blocks])
                        lines = snippet_text.splitlines()

                        # Collect lines from first "gift message" to second (or end)
                        capturing = False
                        captured_lines = []
                        for ln in lines:
                            if re.search(r"gift message", ln, re.IGNORECASE):
                                if not capturing:
                                    capturing = True  # start from this line
                                else:
                                    # second occurrence => stop
                                    capturing = False
                                    break
                            if capturing:
                                captured_lines.append(ln)

                        final_text = "\n".join(captured_lines).strip()

                        # If we extracted something new, handle duplication checks
                        if final_text and do_text_extraction and gift_doc:
                            if final_text not in gift_messages_per_page[page_index]:
                                gift_messages_per_page[page_index].add(final_text)
                                # add to Word doc
                                p = gift_doc.add_paragraph()
                                p.add_run(f"[Page {page_index}] ").bold = True
                                p.add_run(final_text)

                        # If user wants to highlight entire snippet region
                        if do_snippet_highlight and snippet_blocks:
                            # unify bounding rectangle for the snippet
                            bounding_rect = None
                            capturing_blocks = False
                            for (txt, rect) in snippet_blocks:
                                if re.search(r"gift message", txt, re.IGNORECASE):
                                    if not capturing_blocks:
                                        capturing_blocks = True
                                    else:
                                        # second => done
                                        break
                                if capturing_blocks:
                                    if bounding_rect is None:
                                        bounding_rect = rect
                                    else:
                                        bounding_rect |= rect

                            if bounding_rect:
                                big_annot = page.add_highlight_annot(bounding_rect)
                                big_annot.set_colors(stroke=purple_color)
                                big_annot.update()

            # (D) Highlight personalization lines if requested, and optionally extract them
            lines_data = get_lines(page)
            current_item_has_custom = False
            # Convert these to lowercase for consistent matching
            stop_keywords = [
                "size:",
                "word or mssg:",
                "word or message:",
                "morse code:",
                "quantity:",
                "sku:",
                "private notes",
                "scheduled to ship by",
                "note from buyer",
                "do the green thing",
                "reuse this paper to make origami, confetti",
                "or your next to-do list.",
            ]

            idx = 0
            while idx < len(lines_data):
                line_text, line_rect, word_list = lines_data[idx]
                lower_line = line_text.lower()

                # Check if line indicates a "custom" item
                is_descriptor = any(desc in lower_line for desc in item_descriptors)
                if is_descriptor:
                    current_item_has_custom = ("custom" in lower_line)

                if personalization_marker in lower_line and current_item_has_custom:
                    # If the user wants to extract text, write it to doc
                    if do_text_extraction and personalization_doc:
                        p = personalization_doc.add_paragraph()
                        p.add_run(f"[Page {page_index}] ").bold = True
                        p.add_run(line_text)

                    # If user wants to highlight personalization lines
                    if options['highlight_personalization']:
                        for (wx0, wy0, wx1, wy1, wtext) in word_list:
                            if "personalization:" in wtext.lower():
                                continue
                            w_rect = fitz.Rect(wx0, wy0, wx1, wy1)
                            hl_annot = page.add_highlight_annot(w_rect)
                            hl_annot.set_colors(stroke=yellow_color)
                            hl_annot.update()

                    # Continue to subsequent lines until we hit stop keywords or another personalization:
                    next_idx = idx + 1
                    while next_idx < len(lines_data):
                        nl_text, nl_rect, nl_words = lines_data[next_idx]
                        lower_nl_text = nl_text.lower()

                        # If any lowercased stop keyword is in the line, stop
                        if any(kw in lower_nl_text for kw in stop_keywords):
                            break
                        if personalization_marker in lower_nl_text:
                            break

                        if do_text_extraction and personalization_doc:
                            p = personalization_doc.add_paragraph()
                            p.add_run(f"[Page {page_index}] ").bold = True
                            p.add_run(nl_text)

                        if options['highlight_personalization']:
                            for (wx0, wy0, wx1, wy1, wtext) in nl_words:
                                if "personalization:" in wtext.lower():
                                    continue
                                w_rect = fitz.Rect(wx0, wy0, wx1, wy1)
                                hl_annot = page.add_highlight_annot(w_rect)
                                hl_annot.set_colors(stroke=yellow_color)
                                hl_annot.update()

                        next_idx += 1

                    idx = next_idx
                else:
                    idx += 1

            # (E) Highlight quantity >=2 if user wants
            if options['highlight_quantity']:
                lines_data = get_lines(page)
                for (line_text, line_rect, word_list) in lines_data:
                    lower_line = line_text.lower()
                    if "quantity:" in lower_line:
                        quantity_index = None
                        for w_i, (wx0, wy0, wx1, wy1, wtext) in enumerate(word_list):
                            if wtext.lower().startswith("quantity:"):
                                quantity_index = w_i
                                break
                        if quantity_index is not None and (quantity_index + 1) < len(word_list):
                            nx0, ny0, nx1, ny1, next_text = word_list[quantity_index + 1]
                            try:
                                val = int(next_text)
                                if val >= 2:
                                    highlight_rect = fitz.Rect(nx0, ny0, nx1, ny1)
                                    hl_annot = page.add_highlight_annot(highlight_rect)
                                    hl_annot.set_colors(stroke=yellow_color)
                                    hl_annot.update()
                            except:
                                pass

            # (F) Extract "ship to" info (always done so we can place it on page0 if watermark is used)
            lines_data = get_lines(page)
            for i in range(len(lines_data)):
                line_text, line_rect, word_list = lines_data[i]
                if "ship to" in line_text.lower():
                    extracted_name = ""
                    if i + 1 < len(lines_data):
                        next_line_text, _, next_line_words = lines_data[i + 1]
                        if len(next_line_words) >= 2:
                            name_words = [w[4] for w in next_line_words[:2]]
                            extracted_name = " ".join(name_words)
                        else:
                            extracted_name = next_line_text
                        ship_to_info.append((page_index, extracted_name))
                    break

            # (G) Add stamps if user checked "Add Stamps"
            if options['add_stamps']:
                page_text_lower = page.get_text().lower()
                found_gift = any(phrase in page_text_lower for phrase in gift_trigger_phrases)
                stamps_to_insert = []

                if found_gift and stamp_images.get("gift"):
                    stamps_to_insert.append(stamp_images["gift"])
                if "igb" in page_text_lower and stamp_images.get("igb"):
                    stamps_to_insert.append(stamp_images["igb"])
                if "upgrade label to priority - bubble" in page_text_lower and stamp_images.get("bubble"):
                    stamps_to_insert.append(stamp_images["bubble"])
                if "show kd" in page_text_lower and stamp_images.get("show"):
                    stamps_to_insert.append(stamp_images["show"])
                if (("pic kd" in page_text_lower) or ("pic " in page_text_lower)) and stamp_images.get("pic"):
                    stamps_to_insert.append(stamp_images["pic"])
                if "short thins" in page_text_lower and stamp_images.get("short"):
                    stamps_to_insert.append(stamp_images["short"])
                if "hjlm" in page_text_lower and stamp_images.get("hjlm"):
                    stamps_to_insert.append(stamp_images["hjlm"])
                if "priority box" in page_text_lower and stamp_images.get("priority"):
                    stamps_to_insert.append(stamp_images["priority"])
                # New: detect "fedex"
                if "fedex" in page_text_lower and stamp_images.get("fedex"):
                    stamps_to_insert.append(stamp_images["fedex"])

                page_rect = page.rect
                # Smaller stamp height, plus reduced distance between stamps
                stamp_height = 70
                y_offset = 50

                for stamp_path in stamps_to_insert:
                    stamp_rect = fitz.Rect(
                        page_rect.width - 150,
                        y_offset,
                        page_rect.width - 20,
                        y_offset + stamp_height
                    )
                    page.insert_image(stamp_rect, filename=stamp_path, keep_proportion=True)
                    # Increase y_offset by stamp height + 10
                    y_offset += (stamp_height + 10)

            # (H) Insert page number + optional watermark text
            page_rect = page.rect
            page.insert_text(
                (25, 25),
                f"Page {page_index} of {total_pages - 1}",
                fontname="helv",
                fontsize=10,
                color=(0, 0, 0)
            )
            if options['apply_watermark'] and watermark_text.strip():
                page.insert_text(
                    (page_rect.width - 150, 35),
                    watermark_text,
                    fontname="helv",
                    fontsize=10,
                    color=(0, 0, 0)
                )

        # (I) On the first blank page, optionally put the 'ship to' info with watermark text
        if options['apply_watermark'] and watermark_text.strip():
            page0 = doc[0]
            x, y = 50, 50
            for (pnum, name) in ship_to_info:
                page0.insert_text(
                    (x, y),
                    f"{watermark_text}{pnum}: {name}",
                    fontname="helv",
                    fontsize=12,
                    color=(0, 0, 0)
                )
                y += 20

        # Finally, save PDF
        doc.save(output_pdf_path, garbage=4, deflate=True)
        doc.close()

        # Save docx if extraction was enabled
        if options['extract_text'] and gift_doc and personalization_doc:
            gift_doc.save(word_output_path)
            personalization_doc.save(personalization_word_path)

        return f"Processed '{os.path.basename(pdf_path)}' successfully."

    except Exception as e:
        return f"Error processing '{os.path.basename(pdf_path)}': {e}"


# --------------------------------------------------------------------------
# 3) GUI Class
# --------------------------------------------------------------------------
class PDFBatchProcessorGUI(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("PDF Batch Processor")
        self.geometry("750x750")
        self.create_widgets()

        # Dictionary of stamp images (add your new 'fed.png')
        self.stamp_images = {
            "gift": resource_path("gift_stamp.png"),
            "igb": resource_path("igb.png"),
            "bubble": resource_path("bubble.png"),
            "show": resource_path("show.png"),
            "pic": resource_path("pic.png"),
            "short": resource_path("short.png"),
            "hjlm": resource_path("HJLM.png"),
            "priority": resource_path("priority.png"),
            # New FedEx stamp
            "fedex": resource_path("fed.png"),
        }

    def create_widgets(self):
        # Frame: select PDFs
        tk.Label(self, text="Select PDF Files (up to 20):").pack(pady=5)
        select_frame = tk.Frame(self)
        select_frame.pack()
        tk.Button(select_frame, text="Browse PDFs", command=self.browse_pdfs).pack(side=tk.LEFT, padx=5)
        tk.Button(select_frame, text="Clear List", command=self.clear_pdf_list).pack(side=tk.LEFT, padx=5)
        self.pdf_listbox = tk.Listbox(self, width=100, height=8)
        self.pdf_listbox.pack(pady=5)

        # Frame: select output directory
        tk.Label(self, text="Select Output Directory:").pack(pady=5)
        out_frame = tk.Frame(self)
        out_frame.pack()
        self.out_dir_entry = tk.Entry(out_frame, width=80)
        self.out_dir_entry.pack(side=tk.LEFT, padx=5)
        tk.Button(out_frame, text="Browse", command=self.browse_output_dir).pack(side=tk.LEFT, padx=5)

        # Frame: Watermark text
        watermark_frame = tk.Frame(self)
        watermark_frame.pack(pady=(10,5))
        tk.Label(watermark_frame, text="Watermark Text:").pack(side=tk.LEFT, padx=5)
        self.watermark_entry = tk.Entry(watermark_frame, width=30)
        self.watermark_entry.pack(side=tk.LEFT, padx=5)

        # Frame: Checkboxes (options)
        options_frame = tk.LabelFrame(self, text="Select Processing Options")
        options_frame.pack(pady=10, fill='x', padx=10)

        self.crop_pages_var = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text="Crop Pages", variable=self.crop_pages_var).pack(anchor='w')

        self.highlight_gmi_var = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text="Highlight phrase 'gift message included'",
                       variable=self.highlight_gmi_var).pack(anchor='w')

        self.highlight_gift_snippet_var = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text="Highlight entire 'gift message' snippet",
                       variable=self.highlight_gift_snippet_var).pack(anchor='w')

        self.extract_text_var = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text="Extract Gift & Personalization text (create Word docs)",
                       variable=self.extract_text_var).pack(anchor='w')

        self.highlight_personalization_var = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text="Highlight Personalization Lines",
                       variable=self.highlight_personalization_var).pack(anchor='w')

        self.highlight_quantity_var = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text="Highlight quantity >= 2",
                       variable=self.highlight_quantity_var).pack(anchor='w')

        self.add_stamps_var = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text="Add Stamps to pages",
                       variable=self.add_stamps_var).pack(anchor='w')

        self.apply_watermark_var = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text="Apply Watermark text (including 'Ship to' info)",
                       variable=self.apply_watermark_var).pack(anchor='w')

        self.generate_2up_thumbs_var = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text="Generate 2-up and thumbnail PDFs",
                       variable=self.generate_2up_thumbs_var).pack(anchor='w')

        # Process button
        tk.Button(self, text="Process Files", command=self.process_files,
                  bg="green", fg="white", width=20).pack(pady=15)

        # Status text
        self.status_text = tk.Text(self, height=12, width=85)
        self.status_text.pack(pady=5)
        self.status_text.config(state=tk.DISABLED)

    def browse_pdfs(self):
        files = filedialog.askopenfilenames(
            title="Select PDF Files",
            filetypes=[("PDF Files", "*.pdf")]
        )
        if files:
            current_files = list(self.pdf_listbox.get(0, tk.END))
            for f in files:
                if len(current_files) >= 20:
                    messagebox.showwarning("Limit Reached", "You can only select up to 20 PDF files.")
                    break
                if f not in current_files:
                    self.pdf_listbox.insert(tk.END, f)
                    current_files.append(f)

    def clear_pdf_list(self):
        self.pdf_listbox.delete(0, tk.END)

    def browse_output_dir(self):
        directory = filedialog.askdirectory(title="Select Output Directory")
        if directory:
            self.out_dir_entry.delete(0, tk.END)
            self.out_dir_entry.insert(0, directory)

    def log_message(self, msg):
        self.status_text.config(state=tk.NORMAL)
        self.status_text.insert(tk.END, msg + "\n")
        self.status_text.config(state=tk.DISABLED)
        self.status_text.see(tk.END)

    def process_files(self):
        pdf_files = list(self.pdf_listbox.get(0, tk.END))
        out_dir = self.out_dir_entry.get().strip()
        if not pdf_files:
            messagebox.showerror("Error", "No PDF files selected.")
            return
        if not out_dir or not os.path.isdir(out_dir):
            messagebox.showerror("Error", "Please select a valid output directory.")
            return

        watermark_text = self.watermark_entry.get().strip()

        # Gather checkbox options into a dictionary
        options = {
            'crop_pages':                self.crop_pages_var.get(),
            'highlight_gmi':             self.highlight_gmi_var.get(),
            'highlight_gift_snippet':    self.highlight_gift_snippet_var.get(),
            'extract_text':              self.extract_text_var.get(),
            'highlight_personalization': self.highlight_personalization_var.get(),
            'highlight_quantity':        self.highlight_quantity_var.get(),
            'add_stamps':                self.add_stamps_var.get(),
            'apply_watermark':           self.apply_watermark_var.get(),
            'generate_2up_thumbs':       self.generate_2up_thumbs_var.get(),
        }

        self.log_message("Processing started...")

        for pdf in pdf_files:
            base = os.path.splitext(os.path.basename(pdf))[0]
            output_pdf_path = os.path.join(out_dir, base + "_modified.pdf")
            word_output_path = os.path.join(out_dir, base + "_GiftMessages.docx")
            personalization_word_path = os.path.join(out_dir, base + "_Personalizations.docx")

            # 1) Process the main PDF based on checkboxes
            result = process_pdf_file(
                pdf_path=pdf,
                output_pdf_path=output_pdf_path,
                word_output_path=word_output_path,
                personalization_word_path=personalization_word_path,
                stamp_images=self.stamp_images,
                watermark_text=watermark_text,
                options=options
            )
            self.log_message(result)

            # 2) Generate 2-up and thumbnail only if selected
            if options['generate_2up_thumbs']:
                two_up_path = os.path.join(out_dir, base + "_2up.pdf")
                create_two_up_pdf(output_pdf_path, two_up_path)
                self.log_message(f"2-up PDF created for {base}")

                thumb_path = os.path.join(out_dir, base + "_thumb.pdf")
                create_six_page_thumbnail_pdf(output_pdf_path, thumb_path)
                self.log_message(f"Thumbnail PDF created for {base}")

        self.log_message("All files processed!")
        messagebox.showinfo("Done", "Processing complete for all selected files.")


# --------------------------------------------------------------------------
# Main
# --------------------------------------------------------------------------
if __name__ == "__main__":
    app = PDFBatchProcessorGUI()
    app.mainloop()
